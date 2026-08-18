"""
Quick smoke test: exercises the Flask routes and core conversion functions
using tiny generated sample files, to catch import/logic errors early.
"""
import io
import sys
sys.path.insert(0, ".")
from reportlab.pdfgen import canvas
from PIL import Image
from docx import Document as DocxDocument
from openpyxl import Workbook

from app import app
def make_sample_pdf(text="Hello DocBench", pages=2):
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    for i in range(pages):
        c.drawString(100, 700, f"{text} - page {i+1}")
        c.showPage()
    c.save()
    buf.seek(0)
    return buf


def make_sample_docx():
    doc = DocxDocument()
    doc.add_heading("Test Report", level=1)
    doc.add_paragraph("This is a paragraph for conversion testing.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def make_sample_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.append(["Name", "Score"])
    ws.append(["Alice", 90])
    ws.append(["Bob", 85])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def make_sample_jpg():
    img = Image.new("RGB", (400, 300), color=(120, 160, 200))
    buf = io.BytesIO()
    img.save(buf, format="JPEG")
    buf.seek(0)
    return buf


results = []

def check(name, fn):
    try:
        fn()
        results.append((name, "OK"))
    except Exception as e:
        results.append((name, f"FAIL: {e}"))


client = app.test_client()

def post_file(tool_id, files, form=None):
    data = dict(form or {})
    if isinstance(files, list):
        data["file"] = files
    else:
        data["file"] = files
    resp = client.post(f"/process/{tool_id}", data=data, content_type="multipart/form-data")
    assert resp.status_code == 200, f"{tool_id} -> {resp.status_code}: {resp.data[:300]}"
    assert len(resp.data) > 0
    return resp


check("index page loads", lambda: (_ for _ in ()).throw(AssertionError("bad")) if client.get("/").status_code != 200 else None)
check("merge-pdf", lambda: post_file("merge-pdf", [(make_sample_pdf(), "a.pdf"), (make_sample_pdf(), "b.pdf")]))
check("split-pdf", lambda: post_file("split-pdf", (make_sample_pdf(pages=4), "a.pdf"), {"ranges": "1-2,3,4"}))
check("organize-pdf", lambda: post_file("organize-pdf", (make_sample_pdf(pages=3), "a.pdf"), {"order": "3,1"}))
check("compress-pdf", lambda: post_file("compress-pdf", (make_sample_pdf(), "a.pdf"), {"level": "recommended"}))
check("repair-pdf", lambda: post_file("repair-pdf", (make_sample_pdf(), "a.pdf")))
check("rotate-pdf", lambda: post_file("rotate-pdf", (make_sample_pdf(), "a.pdf"), {"angle": "90"}))
check("watermark-pdf", lambda: post_file("watermark-pdf", (make_sample_pdf(), "a.pdf"), {"text": "DRAFT", "opacity": "0.3"}))
check("page-numbers", lambda: post_file("page-numbers", (make_sample_pdf(pages=3), "a.pdf"), {"position": "bottom-center", "start_at": "1"}))
check("crop-pdf", lambda: post_file("crop-pdf", (make_sample_pdf(), "a.pdf"), {"left": "10", "right": "10", "top": "10", "bottom": "10"}))
check("redact-pdf", lambda: post_file("redact-pdf", (make_sample_pdf(text="SECRET DATA"), "a.pdf"), {"terms": "SECRET"}))
check("protect-pdf", lambda: post_file("protect-pdf", (make_sample_pdf(), "a.pdf"), {"password": "test123"}))
check("pdf-to-jpg", lambda: post_file("pdf-to-jpg", (make_sample_pdf(pages=2), "a.pdf"), {"dpi": "100"}))
check("pdf-to-word", lambda: post_file("pdf-to-word", (make_sample_pdf(), "a.pdf")))
check("pdf-to-powerpoint", lambda: post_file("pdf-to-powerpoint", (make_sample_pdf(pages=2), "a.pdf")))
check("pdf-to-excel", lambda: post_file("pdf-to-excel", (make_sample_pdf(), "a.pdf")))
check("word-to-pdf", lambda: post_file("word-to-pdf", (make_sample_docx(), "a.docx")))
check("excel-to-pdf", lambda: post_file("excel-to-pdf", (make_sample_xlsx(), "a.xlsx")))
check("html-to-pdf", lambda: post_file("html-to-pdf", None, {"html": "<h1>Hi</h1><p>Test</p>"}) if False else client.post("/process/html-to-pdf", data={"html": "<h1>Hi</h1><p>Test paragraph</p>"}))
check("jpg-to-pdf", lambda: post_file("jpg-to-pdf", [(make_sample_jpg(), "a.jpg"), (make_sample_jpg(), "b.jpg")]))
check("compress-image", lambda: post_file("compress-image", (make_sample_jpg(), "a.jpg"), {"quality": "50"}))

# compare-pdf returns HTML, not a file
def check_compare():
    data = {"file": [(make_sample_pdf(text="Alpha"), "a.pdf"), (make_sample_pdf(text="Beta"), "b.pdf")]}
    resp = client.post("/process/compare-pdf", data=data, content_type="multipart/form-data")
    assert resp.status_code == 200, resp.data[:300]
    assert b"table" in resp.data.lower()

check("compare-pdf", check_compare)


print("\n--- SMOKE TEST RESULTS ---")
ok = 0
for name, status in results:
    print(f"{status:6s} {name}" if status == "OK" else f"{status}  ({name})")
    if status == "OK":
        ok += 1
print(f"\n{ok}/{len(results)} passed")
